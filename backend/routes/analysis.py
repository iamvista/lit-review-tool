"""
AI 分析 API
Analysis Routes - AI 輔助研究缺口識別
"""

from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from models import db, Project, Paper, GapAnalysis, User
from services.ai_analyzer import AIAnalyzer
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

analysis_bp = Blueprint('analysis', __name__, url_prefix='/api/analysis')


@analysis_bp.route('/projects/<int:project_id>/analyze', methods=['POST'])
@jwt_required()
def analyze_project(project_id):
    """
    執行專案的 AI 分析
    POST /api/analysis/projects/:id/analyze

    Body: {
        "analysis_type": "comprehensive" | "domain_history" | "core_problems" | ...,
        "title": "分析報告標題（可選）"
    }
    """
    user_id = int(get_jwt_identity())

    # 驗證專案所有權
    project = Project.query.filter_by(id=project_id, user_id=user_id).first()
    if not project:
        return jsonify({'error': '專案不存在或無權限'}), 404

    # 獲取請求參數
    data = request.get_json() or {}
    analysis_type = data.get('analysis_type', 'comprehensive')
    title = data.get('title', f'{project.name} - AI 分析報告')

    # 獲取用戶的 API Key
    user = User.query.get(user_id)
    if not user:
        return jsonify({'error': '用戶不存在'}), 404

    api_key = user.get_anthropic_api_key()
    if not api_key:
        # 如果用戶沒有設定，嘗試使用環境變數（向後相容）
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            return jsonify({
                'error': '未配置 AI API 金鑰',
                'message': '請在設定頁面設置您的 Anthropic API Key'
            }), 403

    try:
        # 獲取專案的所有論文
        papers = Paper.query.filter_by(project_id=project_id).order_by(Paper.year).all()

        if not papers:
            return jsonify({'error': '專案中沒有論文'}), 400

        if len(papers) < 3:
            return jsonify({
                'error': '論文數量不足',
                'message': '建議至少有 3 篇論文才能進行有效分析'
            }), 400

        # 創建 AI 分析器
        analyzer = AIAnalyzer(api_key=api_key)

        # 執行分析
        result = analyzer.analyze_with_context(
            papers=papers,
            domain_name=project.domain or project.name,
            analysis_type=analysis_type
        )

        # 保存分析結果
        gap_analysis = GapAnalysis(
            project_id=project_id,
            title=title,
            analysis_type=analysis_type,
            full_text=result.get('full_text') or result.get('content', ''),
            model_used=result.get('model'),
            papers_analyzed=len(papers)
        )

        # 如果是綜合分析，嘗試解析結構化內容
        if analysis_type == 'comprehensive' and 'full_text' in result:
            gap_analysis.summary = result.get('full_text')

        db.session.add(gap_analysis)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'AI 分析完成',
            'analysis': gap_analysis.to_dict()
        }), 201

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'分析失敗: {str(e)}'}), 500


@analysis_bp.route('/projects/<int:project_id>/analyses', methods=['GET'])
@jwt_required()
def get_project_analyses(project_id):
    """
    獲取專案的所有分析報告
    GET /api/analysis/projects/:id/analyses
    """
    user_id = int(get_jwt_identity())

    # 驗證專案所有權
    project = Project.query.filter_by(id=project_id, user_id=user_id).first()
    if not project:
        return jsonify({'error': '專案不存在或無權限'}), 404

    try:
        # 獲取所有分析報告
        analyses = GapAnalysis.query.filter_by(project_id=project_id).order_by(
            GapAnalysis.generated_at.desc()
        ).all()

        return jsonify({
            'success': True,
            'analyses': [a.to_dict() for a in analyses],
            'total': len(analyses)
        }), 200

    except Exception as e:
        return jsonify({'error': f'獲取分析報告失敗: {str(e)}'}), 500


@analysis_bp.route('/analyses/<int:analysis_id>', methods=['GET'])
@jwt_required()
def get_analysis(analysis_id):
    """
    獲取單個分析報告詳情
    GET /api/analysis/analyses/:id
    """
    user_id = int(get_jwt_identity())

    try:
        analysis = GapAnalysis.query.get(analysis_id)
        if not analysis:
            return jsonify({'error': '分析報告不存在'}), 404

        # 驗證專案所有權
        project = Project.query.filter_by(id=analysis.project_id, user_id=user_id).first()
        if not project:
            return jsonify({'error': '無權限訪問此分析報告'}), 403

        return jsonify({
            'success': True,
            'analysis': analysis.to_dict()
        }), 200

    except Exception as e:
        return jsonify({'error': f'獲取分析報告失敗: {str(e)}'}), 500


@analysis_bp.route('/analyses/<int:analysis_id>', methods=['DELETE'])
@jwt_required()
def delete_analysis(analysis_id):
    """
    刪除分析報告
    DELETE /api/analysis/analyses/:id
    """
    user_id = int(get_jwt_identity())

    try:
        analysis = GapAnalysis.query.get(analysis_id)
        if not analysis:
            return jsonify({'error': '分析報告不存在'}), 404

        # 驗證專案所有權
        project = Project.query.filter_by(id=analysis.project_id, user_id=user_id).first()
        if not project:
            return jsonify({'error': '無權限刪除此分析報告'}), 403

        db.session.delete(analysis)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '分析報告已刪除'
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'刪除失敗: {str(e)}'}), 500


@analysis_bp.route('/check-api-key', methods=['GET'])
@jwt_required()
def check_api_key():
    """
    檢查 AI API 金鑰是否已配置
    GET /api/analysis/check-api-key
    """
    user_id = int(get_jwt_identity())
    user = User.query.get(user_id)

    if not user:
        return jsonify({'error': '用戶不存在'}), 404

    # 優先檢查用戶的 API Key
    user_api_key = user.get_anthropic_api_key()
    # 如果用戶沒有，檢查環境變數（向後相容）
    env_api_key = os.environ.get('ANTHROPIC_API_KEY')

    has_key = bool(user_api_key or env_api_key)
    source = 'user' if user_api_key else ('env' if env_api_key else 'none')

    return jsonify({
        'configured': has_key,
        'source': source,
        'message': 'AI API 金鑰已配置' if has_key else '尚未配置 AI API 金鑰，請前往設定頁面設置'
    }), 200


@analysis_bp.route('/analyses/<int:analysis_id>/export/docx', methods=['POST'])
@jwt_required()
def export_analysis_docx(analysis_id):
    """
    匯出分析報告為 DOCX 格式
    POST /api/analysis/analyses/{analysis_id}/export/docx
    """
    user_id = int(get_jwt_identity())

    # 獲取分析報告
    analysis = GapAnalysis.query.get(analysis_id)
    if not analysis:
        return jsonify({'error': '分析報告不存在'}), 404

    # 驗證權限
    project = Project.query.get(analysis.project_id)
    if not project or project.user_id != user_id:
        return jsonify({'error': '無權訪問此分析報告'}), 403

    try:
        # 創建 Word 文檔
        doc = Document()

        # 設置文檔標題
        title = doc.add_heading(analysis.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元數據
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"生成時間：{analysis.generated_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
        metadata_para.add_run(f"分析論文：{analysis.papers_analyzed} 篇\n")
        metadata_para.add_run(f"使用模型：{analysis.model_used or 'Claude'}\n")
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 解析 Markdown 並轉換為 Word 格式
        content = analysis.full_text or ''
        lines = content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                doc.add_paragraph()
                i += 1
                continue

            # H1 標題 (## )
            if line.startswith('## '):
                heading_text = line.replace('## ', '')
                heading = doc.add_heading(heading_text, level=1)
                heading.runs[0].font.size = Pt(18)
                heading.runs[0].font.bold = True

            # H2 標題 (### )
            elif line.startswith('### '):
                heading_text = line.replace('### ', '')
                heading = doc.add_heading(heading_text, level=2)
                heading.runs[0].font.size = Pt(14)
                heading.runs[0].font.bold = True

            # H3 標題 (#### )
            elif line.startswith('#### '):
                heading_text = line.replace('#### ', '')
                heading = doc.add_heading(heading_text, level=3)
                heading.runs[0].font.size = Pt(12)
                heading.runs[0].font.bold = True

            # 無序列表 (- 開頭)
            elif line.startswith('- ') or line.startswith('* '):
                list_text = line[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(para, list_text)

            # 有序列表 (數字. 開頭)
            elif re.match(r'^\d+\.\s', line):
                list_text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(style='List Number')
                _add_formatted_text(para, list_text)

            # 普通段落
            else:
                para = doc.add_paragraph()
                _add_formatted_text(para, line)

            i += 1

        # 保存到內存
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)

        # 返回文件
        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{analysis.title}.docx"
        )

    except Exception as e:
        print(f"匯出 DOCX 失敗: {e}")
        return jsonify({'error': f'匯出失敗：{str(e)}'}), 500


def _add_formatted_text(paragraph, text):
    """
    添加格式化文本到段落（處理粗體、斜體等）
    """
    # 處理粗體 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗體文本
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.font.bold = True
        elif part:
            # 普通文本
            paragraph.add_run(part)
